from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from typing import Dict, Any, List, Union
import os
import logging

# Setup logger for this module
logger = logging.getLogger(__name__)

def parse_docx(file_path: Union[str, os.PathLike]) -> Dict[str, Any]:
    """
    Parses a .docx file and extracts text from paragraphs, tables, and core properties.
    Includes detailed logging for the extraction process.

    Args:
        file_path: Path to the .docx file.

    Returns:
        A dictionary containing extracted data.
    """
    logger.info(f"--- Starting DOCX parsing for: {file_path} ---")
    extracted_data = {
        "paragraphs": [],
        "tables": [],
        "metadata": {},
        "document_text": ""
    }
    total_chars_extracted = 0

    try:
        document = Document(file_path)
        logger.info(f"Successfully opened document: {file_path}")
        
        # Extract core properties (metadata)
        cp = document.core_properties
        meta = {
            "author": cp.author, "category": cp.category, "comments": cp.comments,
            "content_status": cp.content_status, "created": cp.created, "identifier": cp.identifier,
            "keywords": cp.keywords, "language": cp.language, "last_modified_by": cp.last_modified_by,
            "last_printed": cp.last_printed, "modified": cp.modified, "revision": cp.revision,
            "subject": cp.subject, "title": cp.title, "version": cp.version
        }
        extracted_data["metadata"] = {k: str(v) if v is not None else "" for k, v in meta.items()} # Ensure strings
        logger.info("Extracted Core Properties (Metadata):")
        for key, value in extracted_data["metadata"].items():
            if value: # Log only if metadata has a value
                logger.debug(f"  - {key.title()}: {str(value)[:100]}") # Log first 100 chars

        all_text_parts = []
        paragraphs_found_with_text = 0
        tables_found_with_text = 0

        # Extract text from paragraphs
        logger.info(f"Found {len(document.paragraphs)} paragraphs in total.")
        for i, para in enumerate(document.paragraphs):
            para_text = para.text.strip()
            if para_text:
                logger.debug(f"  Para {i+1}: Extracted '{para_text[:100]}...' (Length: {len(para_text)})")
                extracted_data["paragraphs"].append(para_text)
                all_text_parts.append(para_text)
                total_chars_extracted += len(para_text)
                paragraphs_found_with_text += 1
            else:
                logger.debug(f"  Para {i+1}: No text content.")
        
        if paragraphs_found_with_text == 0 and len(document.paragraphs) > 0:
            logger.warning("⚠️ No text extracted from any paragraphs, though paragraphs were present.")
        elif len(document.paragraphs) == 0:
            logger.warning("⚠️ No paragraphs found in the document.")


        # Extract text from tables
        logger.info(f"Found {len(document.tables)} tables in total.")
        for i, table in enumerate(document.tables):
            table_data = {"table_index": i, "rows": []}
            table_text_parts_for_doc = [] # For concatenated document_text
            current_table_has_text = False
            logger.debug(f"  Table {i+1}: Processing {len(table.rows)} rows.")
            for r_idx, row in enumerate(table.rows):
                row_texts = []
                cell_texts_for_concat = []
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        logger.debug(f"    Table {i+1}, Row {r_idx+1}, Cell {c_idx+1}: Extracted '{cell_text[:50]}...' (Length: {len(cell_text)})")
                        row_texts.append(cell_text)
                        cell_texts_for_concat.append(cell_text)
                        total_chars_extracted += len(cell_text)
                        current_table_has_text = True
                    else:
                        logger.debug(f"    Table {i+1}, Row {r_idx+1}, Cell {c_idx+1}: No text content.")
                        cell_texts_for_concat.append("") # Keep structure for | join
                if row_texts: # If any cell in the row had text
                    table_data["rows"].append(row_texts)
                table_text_parts_for_doc.append(" | ".join(cell_texts_for_concat))
            
            if current_table_has_text:
                tables_found_with_text +=1
            
            extracted_data["tables"].append(table_data)
            all_text_parts.append("\\n".join(table_text_parts_for_doc)) # Add table text to main text list
        
        if tables_found_with_text == 0 and len(document.tables) > 0:
            logger.warning("⚠️ No text extracted from any tables, though tables were present.")
        elif len(document.tables) == 0:
            logger.info("ℹ️ No tables found in the document.")
        
        extracted_data["document_text"] = "\\n\\n".join(all_text_parts).strip() # Concatenate all text parts

        if not extracted_data["document_text"]:
            logger.warning(f"⚠️ No text content could be extracted overall from {file_path}.")
        else:
            logger.info(f"✅ Successfully extracted {total_chars_extracted} characters in total from {file_path}.")
            logger.debug(f"Final 'document_text' preview (first 200 chars): {extracted_data['document_text'][:200]}...")

        logger.info(f"--- Finished DOCX parsing for: {file_path} ---")
        return extracted_data

    except PackageNotFoundError:
        logger.error(f"❌ Error: '{file_path}' is not a valid DOCX file or is corrupted.", exc_info=True)
        return {"error": f"'{file_path}' is not a valid DOCX file or is corrupted."}
    except Exception as e:
        logger.error(f"❌ An unexpected error occurred while parsing DOCX file {file_path}: {e}", exc_info=True)
        return {"error": f"An unexpected error occurred while parsing {file_path}: {str(e)}"}

if __name__ == '__main__':
    # Configure basic logging for direct script execution
    if not logging.getLogger().handlers: # Avoid adding multiple handlers if script is re-run in some envs
        logging.basicConfig(level=logging.DEBUG, 
                            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                            handlers=[logging.StreamHandler()]) # Ensure logs go to console

    example_file_name = "example_parser_test.docx" 
    # Ensure the test file has a unique name to avoid conflict with other example.docx files

    # Create a dummy DOCX file for testing if python-docx is available
    if not os.path.exists(example_file_name):
        try:
            doc = Document()
            doc.add_heading('Document Title for Test', 0)
            doc.add_paragraph('This is a test paragraph with some bold text and also some regular text.')
            doc.add_paragraph('Another paragraph with more information, perhaps a list:')
            doc.add_paragraph('Item 1: First item', style='ListBullet')
            doc.add_paragraph('Item 2: Second item', style='ListBullet')
            
            # Add a table
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Header A'
            table.cell(0, 1).text = 'Header B'
            table.cell(1, 0).text = 'Data 1'
            table.cell(1, 1).text = 'Data 2 with some more text to make it longer.'
            doc.add_paragraph('A paragraph after the table.')

            doc.core_properties.author = "Test Author - DOCX Parser"
            doc.core_properties.title = "Sample Document for DOCX Parser Test"
            doc.save(example_file_name)
            logger.info(f"Created dummy '{example_file_name}' for testing.")
        except ImportError:
            logger.warning("python-docx library not found. Skipping dummy DOCX creation.")
            logger.info(f"Please create or place an '{example_file_name}' in the directory '{os.getcwd()}' or modify the script.")
        except Exception as e:
            logger.error(f"Error creating dummy DOCX '{example_file_name}': {e}", exc_info=True)

    if os.path.exists(example_file_name):
        logger.info(f"--- Running parsing test for '{example_file_name}' ---")
        parsed_data = parse_docx(example_file_name)

        if parsed_data.get("error"):
            logger.error(f"Error during parsing test: {parsed_data['error']}")
        else:
            # logger.debug(f"Overall Document Text (Concatenated):\\n{parsed_data.get('document_text', '')[:500]}...")
            # logger.debug("Metadata:")
            # for key, value in parsed_data.get("metadata", {}).items():
            #     logger.debug(f"  {key.title()}: {value}")
            # logger.debug("Paragraphs (first 5):")
            # for i, p_text in enumerate(parsed_data.get("paragraphs", [])[:5]):
            #     logger.debug(f"  Para {i+1}: {p_text[:100]}...")
            # if len(parsed_data.get("paragraphs", [])) > 5: logger.debug("  ...")
            # logger.debug("Tables (first 2):")
            # for t_data in parsed_data.get("tables", [])[:2]:
            #     logger.debug(f"  Table {t_data.get('table_index', '')}:")
            #     for r_idx, row in enumerate(t_data.get("rows", [])[:3]):
            #         logger.debug(f"    Row {r_idx+1}: {row}")
            #     if len(t_data.get("rows", [])) > 3: logger.debug("    ...")
            # if len(parsed_data.get("tables", [])) > 2: logger.debug("  ...")
            # The detailed logging is now inside parse_docx, so just a summary here is fine.
            logger.info(f"Test parsing of '{example_file_name}' complete.")
            logger.info(f"Total characters extracted by test: {len(parsed_data.get('document_text', ''))}")
            
        logger.info(f"--- Finished parsing test for '{example_file_name}' ---")
    else:
        logger.warning(f"Skipping parsing test as '{example_file_name}' was not found and dummy creation failed/skipped.") 